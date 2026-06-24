"""Unit tests for the SOW feature: template config, estimate mapper, agent fallback,
client-fact token resolution, and the python-docx renderer.

No test makes a real LLM call: the agent boundary is monkeypatched (or the no-key fallback
path is exercised), so these run offline and deterministically.
"""

from __future__ import annotations

import io

import pytest
from docx import Document

from sow import composer as composer_mod
from sow.composer import _normalize_client_refs, build_sow_document
from sow.config import SowTemplateError, clear_template_cache, load_sow_template
from sow.mapper import RENDERERS, fee_table, resource_summary, schedule_table, total_investment
from sow.models import SowClientFacts, SowDocument, SowSectionContent, SowTable
from sow.renderer import render_docx
from sow.vendor_guard import generalize_vendor_tech

from ._sow_fixtures import AI_FEE_TOTAL, MANUAL_FEE_TOTAL, make_completed_envelope


# --- template config ----------------------------------------------------------------
def test_baseline_template_loads_and_validates() -> None:
    t = load_sow_template()
    assert t.id == "default_sow"
    assert len(t.sections) >= 10
    # Every estimate/hybrid section names a renderer that exists in the registry.
    for s in t.sections:
        if s.source == "estimate":
            assert s.renderer in RENDERERS, s.id
        if s.source == "hybrid" and s.renderer:
            assert s.renderer in RENDERERS, s.id
    # Branding + the canonical placeholders are present.
    assert "Your Success is our success" in t.branding.commitment
    assert "[CLIENT NAME]" in {p.token for p in t.placeholders}


def test_unknown_template_raises() -> None:
    with pytest.raises(SowTemplateError):
        load_sow_template("does_not_exist")


def test_company_name_is_not_hardcoded_required_from_template() -> None:
    """The delivering firm's name has no code default — it must come from the template YAML."""
    from pydantic import ValidationError

    from sow.models import SowBranding

    with pytest.raises(ValidationError):
        SowBranding()  # `company` is required (no hardcoded firm-name default)
    assert SowBranding(company="Acme").company == "Acme"


def test_sow_company_name_env_overrides_template(monkeypatch) -> None:
    """SOW_COMPANY_NAME injects the firm's name at deploy time, overriding the YAML — so no
    specific firm is committed to the repo."""
    from config import get_settings

    monkeypatch.setenv("SOW_COMPANY_NAME", "Globex Corporation")
    get_settings.cache_clear()
    clear_template_cache()
    try:
        assert load_sow_template().branding.company == "Globex Corporation"
    finally:
        # Restore clean cached state for the rest of the suite (env is auto-restored).
        get_settings.cache_clear()
        clear_template_cache()


def test_company_override_is_read_live_after_config_reload(monkeypatch) -> None:
    """Regression: the company override must be resolved LIVE (via the ``config`` module), not
    frozen into the parse cache nor read through a reference captured at import time.

    A prior test reloading ``config`` (as ``test_langfuse_wrapper`` does) used to leave
    ``sow.config`` reading a stale ``get_settings`` whose cached ``Settings`` predated the env —
    so the override was silently skipped. This pins the live read so that can't regress."""
    import importlib

    import config

    # Simulate the cross-test pollution: reload config so any import-time-captured reference
    # would now be stale. The override must still follow the *current* config.get_settings.
    importlib.reload(config)
    monkeypatch.setenv("SOW_COMPANY_NAME", "Initech LLC")
    config.get_settings.cache_clear()
    try:
        # NOTE: deliberately do NOT clear the template parse cache — proving the override is
        # applied on top of the (possibly cached) base, not baked into it.
        assert load_sow_template().branding.company == "Initech LLC"
    finally:
        config.get_settings.cache_clear()
        clear_template_cache()


def test_company_override_does_not_mutate_cached_base(monkeypatch) -> None:
    """The override is applied to a copy: loading with SOW_COMPANY_NAME set must not poison the
    cached base for a later load without it (the bug that made loads order-dependent)."""
    from config import get_settings

    clear_template_cache()
    get_settings.cache_clear()
    baseline = load_sow_template().branding.company  # YAML default, no override
    try:
        monkeypatch.setenv("SOW_COMPANY_NAME", "Umbrella Corp")
        get_settings.cache_clear()
        assert load_sow_template().branding.company == "Umbrella Corp"

        # Drop the env and the settings cache; the base must be pristine again (not "Umbrella").
        monkeypatch.delenv("SOW_COMPANY_NAME", raising=False)
        get_settings.cache_clear()
        assert load_sow_template().branding.company == baseline
    finally:
        get_settings.cache_clear()
        clear_template_cache()


async def test_company_token_resolves_and_is_swappable(monkeypatch) -> None:
    """[COMPANY] resolves from branding.company everywhere (sections + footer); changing that
    one value re-brands the whole SOW — proving the name is dependency-injected, not hardcoded."""
    base = load_sow_template()
    original = base.branding.company
    assert original  # the single source, loaded from the YAML / env

    swapped = base.model_copy(deep=True)
    swapped.branding.company = "Acme Consulting"

    async def _gen(template, envelope, scenario):
        prose = {s.id: "[COMPANY] will deliver the work." for s in template.sections if s.source == "llm"}
        return prose, SowClientFacts()

    monkeypatch.setattr(composer_mod, "load_sow_template", lambda *a, **k: swapped)
    monkeypatch.setattr(composer_mod, "generate_prose", _gen)

    env = make_completed_envelope()
    doc = await build_sow_document(env, "ai_assisted")
    parts = [s.text for s in doc.sections] + [b for s in doc.sections for b in s.bullets]
    parts += [c for s in doc.sections if s.table for row in s.table.rows for c in row]
    parts += [sig.party for s in doc.sections for sig in s.signatories]
    blob = "\n".join(parts)
    assert "[COMPANY]" not in blob  # token fully resolved
    assert "Acme Consulting" in blob  # the injected name appears
    assert original not in blob  # the pre-swap name is fully replaced — nothing hardcoded

    # The footer renders [COMPANY] too (it's outside the composer's token pass).
    rendered = Document(io.BytesIO(render_docx(doc, swapped)))
    footer = rendered.sections[0].footer.paragraphs[0].text
    assert "Acme Consulting" in footer
    assert original not in footer


def test_no_checkbox_or_selection_is_preselected() -> None:
    """Per requirement: no checkbox/selection field is pre-checked — left to the user.
    A checked box is "☒"; empty boxes are "☐"."""
    template = load_sow_template()
    for s in template.sections:
        assert "☒" not in (s.text or ""), f"section {s.id} has a pre-checked box"
        for b in s.bullets:
            assert "☒" not in b, f"section {s.id} bullet has a pre-checked box"
    # The Expenses section still offers the (empty) Yes/No boxes for the user to choose.
    expenses = next(s for s in template.sections if s.id == "expenses")
    assert expenses.text.count("☐") >= 2


def test_template_includes_legal_sections() -> None:
    """The Agreement Parameters + Approval legal sections live in the YAML config."""
    by_id = {s.id: s for s in load_sow_template().sections}
    assert "agreement_parameters" in by_id
    assert "approval" in by_id
    ap = by_id["agreement_parameters"]
    assert ap.source == "boilerplate" and ap.style == "paragraph"
    assert "complete agreement" in ap.text
    approval = by_id["approval"]
    assert "non-disclosure agreement" in approval.text
    assert "purchase order" in approval.text
    assert "\n\n" in approval.text  # multi-paragraph


async def test_legal_sections_render_as_separate_paragraphs(monkeypatch) -> None:
    async def _gen(template, envelope, scenario):
        return {s.id: "x" for s in template.sections if s.source == "llm"}, SowClientFacts()

    monkeypatch.setattr(composer_mod, "generate_prose", _gen)
    env = make_completed_envelope()
    doc = await build_sow_document(env, "ai_assisted")
    reopened = Document(io.BytesIO(render_docx(doc, load_sow_template())))
    paras = [p.text for p in reopened.paragraphs]
    blob = "\n".join(paras)
    assert "Agreement Parameters" in blob and "Approval" in blob
    # The Approval section's paragraphs render as DISTINCT paragraphs, not one merged blob.
    purchase_para = next(p for p in paras if "purchase order" in p)
    assert "non-disclosure" not in purchase_para


# --- estimate mapper ----------------------------------------------------------------
def test_fee_table_rows_sum_to_total_and_follow_scenario() -> None:
    env = make_completed_envelope()

    ai = fee_table(env, "ai_assisted")
    assert ai.columns == ["Resource", "Hourly Rate", "Estimated Hours", "Estimated Investment"]
    # One row per staffed role + a Total row.
    assert len(ai.rows) == len(env.final_estimate.headcount_by_role) + 1
    assert ai.rows[-1][0] == "Total"
    assert ai.rows[-1][-1] == f"${AI_FEE_TOTAL:,.0f}"
    assert total_investment(env, "ai_assisted") == AI_FEE_TOTAL

    manual = fee_table(env, "manual_only")
    assert manual.rows[-1][-1] == f"${MANUAL_FEE_TOTAL:,.0f}"
    assert total_investment(env, "manual_only") == MANUAL_FEE_TOTAL
    # The scenario actually changes the numbers.
    assert ai.rows[-1][-1] != manual.rows[-1][-1]


def test_schedule_table_has_a_row_per_phase() -> None:
    env = make_completed_envelope()
    table = schedule_table(env, "ai_assisted")
    assert len(table.rows) == len(env.final_estimate.phases)
    assert table.rows[0][0] == "Discovery & Analysis"


def test_resource_summary_quotes_the_fee_total() -> None:
    env = make_completed_envelope()
    text = resource_summary(env, "ai_assisted")
    assert f"${AI_FEE_TOTAL:,.0f}" in text
    assert "time-and-materials" in text


# --- agent fallback (no LLM) --------------------------------------------------------
async def test_build_document_falls_back_without_llm(monkeypatch) -> None:
    """When the LLM call raises (no API key), every llm section is still filled and every
    client token degrades to a placeholder (never a guessed value)."""

    async def _boom(**kwargs):
        raise RuntimeError("ANTHROPIC_API_KEY is not set")

    # Force the agent's LLM boundary to fail so generate_prose uses the deterministic path.
    monkeypatch.setattr("sow.agent.call_structured", _boom)

    env = make_completed_envelope()
    doc = await build_sow_document(env, "ai_assisted")

    by_id = {s.id: s for s in doc.sections}
    # Prose sections are non-empty.
    assert by_id["services"].text.strip()
    assert by_id["deliverables"].bullets
    # The fallback prose references the client via the token, never "the Client".
    assert "the Client" not in by_id["services"].text
    assert "[CLIENT NAME]" in by_id["project"].text
    # Client tokens unresolved → reported as placeholders and literal in the cover.
    assert "[CLIENT NAME]" in doc.placeholders
    assert "[CLIENT NAME]" in by_id["cover"].text
    # Estimate-derived tokens ARE resolved (not placeholders).
    assert "[TOTAL INVESTMENT]" not in doc.placeholders
    # The fee table is populated from the estimate.
    assert by_id["fees"].table is not None
    assert by_id["fees"].table.rows[-1][-1] == f"${AI_FEE_TOTAL:,.0f}"


async def test_fallback_prose_comes_from_template_config(monkeypatch) -> None:
    """With no LLM, each llm section's prose is driven by the template's `fallback` /
    `fallback_item` config — not hardcoded in code."""

    async def _boom(**kwargs):
        raise RuntimeError("ANTHROPIC_API_KEY is not set")

    monkeypatch.setattr("sow.agent.call_structured", _boom)

    services_fallback = next(s for s in load_sow_template().sections if s.id == "services").fallback
    assert "blended team of local consultants" in services_fallback  # the config carries it

    env = make_completed_envelope()
    doc = await build_sow_document(env, "ai_assisted")
    by_id = {s.id: s for s in doc.sections}
    # The configured fallback fragment + the {name} substitution survive into the prose.
    assert "blended team of local consultants" in by_id["services"].text
    assert env.project_name in by_id["services"].text  # {name} substituted
    # Bullet sections expand `fallback_item` once per phase.
    assert len(by_id["deliverables"].bullets) == len(env.final_estimate.phases)
    assert "deliverables" in by_id["deliverables"].bullets[0].lower()


# --- client-fact extraction → token resolution --------------------------------------
async def test_extracted_client_name_substituted_else_placeholder(monkeypatch) -> None:
    """Extract-when-present: a grounded client_name fills [CLIENT NAME] (and drops it from
    placeholders); facts left null stay literal placeholders."""

    async def _fake_generate(template, envelope, scenario):
        prose = {s.id: f"{s.id} prose" for s in template.sections if s.source == "llm"}
        return prose, SowClientFacts(client_name="Acme Bank Corporation")  # sow_number etc. null

    monkeypatch.setattr(composer_mod, "generate_prose", _fake_generate)

    env = make_completed_envelope()
    doc = await build_sow_document(env, "ai_assisted")
    cover = next(s for s in doc.sections if s.id == "cover").text

    assert "Acme Bank Corporation" in cover
    assert "[CLIENT NAME]" not in cover
    assert "[CLIENT NAME]" not in doc.placeholders
    # Un-extracted client facts remain placeholders.
    assert "[SOW NUMBER]" in doc.placeholders
    assert "[MSA DATE]" in doc.placeholders


def test_normalize_client_refs_rewrites_the_client_phrasing() -> None:
    assert _normalize_client_refs("We help the Client succeed.") == "We help [CLIENT NAME] succeed."
    assert _normalize_client_refs("the Client's goals") == "[CLIENT NAME]'s goals"
    assert _normalize_client_refs("Deliver for the client.") == "Deliver for [CLIENT NAME]."
    # Must NOT touch unrelated technical phrasing.
    assert _normalize_client_refs("the client-side rendering layer") == "the client-side rendering layer"


def test_generalize_vendor_tech_replaces_brands_with_capabilities() -> None:
    assert "Fargate" not in generalize_vendor_tech("Deploy to AWS ECS Fargate.")
    assert "managed container orchestration" in generalize_vendor_tech("Deploy to AWS ECS Fargate.")
    assert generalize_vendor_tech("set up a RabbitMQ queue") == "set up a message queuing queue"
    assert "Kubernetes" not in generalize_vendor_tech("Run on Kubernetes")
    assert "Snowflake" not in generalize_vendor_tech("Load into Snowflake")
    assert "Auth0" not in generalize_vendor_tech("Integrate Auth0 for login")
    # bare cloud provider
    assert "AWS" not in generalize_vendor_tech("Hosted on AWS")
    # Unrelated prose is untouched.
    assert generalize_vendor_tech("a REST API with OAuth") == "a REST API with OAuth"


async def test_vendor_names_scrubbed_from_built_document(monkeypatch) -> None:
    """A brand that leaks into agent prose is generalized in the final SOW (deterministic guard)."""

    async def _fake_generate(template, envelope, scenario):
        prose = {
            s.id: "[COMPANY] will deploy to AWS ECS Fargate with a RabbitMQ message bus."
            for s in template.sections
            if s.source == "llm"
        }
        return prose, SowClientFacts()

    monkeypatch.setattr(composer_mod, "generate_prose", _fake_generate)
    env = make_completed_envelope()
    doc = await build_sow_document(env, "ai_assisted")
    blob = "\n".join([s.text for s in doc.sections] + [b for s in doc.sections for b in s.bullets])
    for brand in ("ECS Fargate", "Fargate", "RabbitMQ"):
        assert brand not in blob, brand
    assert "managed container orchestration" in blob


async def test_the_client_in_llm_prose_is_rewritten_then_resolved(monkeypatch) -> None:
    """Agent prose saying "the Client" is forced to the token, which then resolves to the
    extracted name (or stays a placeholder) — the SOW never shows the generic "the Client"."""

    async def _fake_generate(template, envelope, scenario):
        prose = {
            s.id: "[COMPANY] will help the Client and the Client's team."
            for s in template.sections
            if s.source == "llm"
        }
        return prose, SowClientFacts(client_name="Acme Bank Corporation")

    monkeypatch.setattr(composer_mod, "generate_prose", _fake_generate)
    env = make_completed_envelope()
    doc = await build_sow_document(env, "ai_assisted")
    services = next(s for s in doc.sections if s.id == "services").text

    assert "the Client" not in services
    assert "the client" not in services
    assert "Acme Bank Corporation" in services  # token resolved to the extracted name
    assert "[CLIENT NAME]" not in doc.placeholders


async def test_all_null_facts_leave_every_client_token(monkeypatch) -> None:
    async def _fake_generate(template, envelope, scenario):
        prose = {s.id: "x" for s in template.sections if s.source == "llm"}
        return prose, SowClientFacts()  # everything null

    monkeypatch.setattr(composer_mod, "generate_prose", _fake_generate)
    env = make_completed_envelope()
    doc = await build_sow_document(env, "ai_assisted")
    assert "[CLIENT NAME]" in doc.placeholders
    assert "[CLIENT REPRESENTATIVE]" in doc.placeholders


# --- renderer (docx round-trips) ----------------------------------------------------
def test_render_docx_reopens_with_headings_total_and_placeholder() -> None:
    template = load_sow_template()
    document = SowDocument(
        estimate_id="e1",
        template_id="default_sow",
        title="STATEMENT OF WORK",
        project_name="Patient Portal Modernization",
        sections=[
            SowSectionContent(id="cover", heading="", kind="cover", text="SOW for [CLIENT NAME]."),
            SowSectionContent(
                id="services", heading="2. Services", kind="paragraph", text="The team will deliver."
            ),
            SowSectionContent(
                id="deliverables",
                heading="Deliverables",
                kind="bullets",
                bullets=["Design artifacts – the wireframes", "Implemented features – the app"],
            ),
            SowSectionContent(
                id="fees",
                heading="5. Fees",
                kind="table",
                table=SowTable(
                    columns=["Resource", "Hourly Rate", "Estimated Hours", "Estimated Investment"],
                    rows=[["Total", "", "1,200", "$236,000"]],
                ),
            ),
        ],
        placeholders=["[CLIENT NAME]"],
    )
    data = render_docx(document, template)
    assert data[:2] == b"PK"  # .docx is a zip

    reopened = Document(io.BytesIO(data))
    all_text = "\n".join(p.text for p in reopened.paragraphs)
    table_text = "\n".join(
        c.text for tbl in reopened.tables for row in tbl.rows for c in row.cells
    )
    assert "STATEMENT OF WORK" in all_text
    assert "5. Fees" in all_text
    assert "Design artifacts – the wireframes" in all_text
    assert "[CLIENT NAME]" in all_text  # placeholders survive verbatim
    assert "$236,000" in table_text


async def test_render_docx_works_for_wbs_method(monkeypatch) -> None:
    """A WBS estimate (no per-phase algorithm) renders identically — SOW reads final_estimate."""

    async def _fake_generate(template, envelope, scenario):
        return {s.id: "x" for s in template.sections if s.source == "llm"}, SowClientFacts()

    monkeypatch.setattr(composer_mod, "generate_prose", _fake_generate)
    env = make_completed_envelope(method="wbs")
    doc = await build_sow_document(env, "ai_assisted")
    data = render_docx(doc, load_sow_template())
    assert data[:2] == b"PK"
