"""Render a resolved ``SowDocument`` to ``.docx`` bytes via python-docx.

Pure presentation: it faithfully emits whatever content the (possibly user-edited)
``SowDocument`` carries — placeholders included — applying the template's branding + style.
No LLM, no estimate access. Dispatches on each section's ``kind``.
"""

from __future__ import annotations

import io
import re

from docx import Document
from docx.document import Document as DocumentT  # the class, for type annotations
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from .models import SowDocument, SowSectionContent, SowTemplate

_PARAGRAPH_SPLIT = re.compile(r"\n\s*\n")


def _rgb(hex_color: str) -> RGBColor:
    try:
        return RGBColor.from_string(hex_color.lstrip("#"))
    except (ValueError, AttributeError):
        return RGBColor.from_string("1F3864")


def _paragraphs(text: str) -> list[str]:
    """Split into display paragraphs (blank line = break); collapse stray single newlines."""
    blocks = _PARAGRAPH_SPLIT.split(text.strip())
    return [" ".join(b.split()) for b in blocks if b.strip()]


def _add_heading(doc: DocumentT, text: str, *, rgb: RGBColor, size: float) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = rgb
    run.font.size = Pt(size)


def _render_table(doc: DocumentT, content: SowSectionContent) -> None:
    table = content.table
    if table is None or not table.columns:
        return
    t = doc.add_table(rows=1, cols=len(table.columns))
    t.style = "Table Grid"
    for cell, col in zip(t.rows[0].cells, table.columns, strict=False):
        cell.text = col
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for row in table.rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            if i < len(cells):
                cells[i].text = str(value)


def _render_signature_block(doc: DocumentT, content: SowSectionContent, *, base_size: float) -> None:
    for para in _paragraphs(content.text):
        doc.add_paragraph(para)
    if not content.signatories:
        return
    doc.add_paragraph()
    t = doc.add_table(rows=1, cols=len(content.signatories))
    for cell, sig in zip(t.rows[0].cells, content.signatories, strict=False):
        head = cell.paragraphs[0]
        run = head.add_run(sig.party)
        run.bold = True
        for field in sig.fields:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.add_run(f"{field}: ").bold = True
            p.add_run("_" * 24)


def _render_section(doc: DocumentT, content: SowSectionContent, *, style: SowStyleSize) -> None:
    if content.heading:
        _add_heading(doc, content.heading, rgb=style.heading_rgb, size=style.heading_size)
    if content.kind == "table":
        _render_table(doc, content)
    elif content.kind == "signature_block":
        _render_signature_block(doc, content, base_size=style.base_size)
    elif content.kind == "bullets":
        for bullet in content.bullets:
            if bullet.strip():
                doc.add_paragraph(bullet, style="List Bullet")
    else:  # paragraph, cover
        for para in _paragraphs(content.text):
            doc.add_paragraph(para)


class SowStyleSize:
    """Small bundle of resolved style values threaded through rendering."""

    def __init__(self, *, heading_rgb: RGBColor, heading_size: float, base_size: float) -> None:
        self.heading_rgb = heading_rgb
        self.heading_size = heading_size
        self.base_size = base_size


def render_docx(document: SowDocument, template: SowTemplate) -> bytes:
    """Render the resolved SOW into ``.docx`` bytes."""
    style = template.style
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = style.font
    normal.font.size = Pt(style.font_size_pt)
    for sect in doc.sections:
        sect.top_margin = sect.bottom_margin = Inches(style.margin_inch)
        sect.left_margin = sect.right_margin = Inches(style.margin_inch)

    heading_rgb = _rgb(style.heading_color)
    size_bundle = SowStyleSize(
        heading_rgb=heading_rgb,
        heading_size=style.font_size_pt + 3,
        base_size=style.font_size_pt,
    )

    # Title + branding lockup.
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(document.title)
    title_run.bold = True
    title_run.font.size = Pt(style.font_size_pt + 8)
    title_run.font.color.rgb = heading_rgb

    if template.branding.tagline:
        tag_p = doc.add_paragraph()
        tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tag_run = tag_p.add_run(f"{template.branding.company}  —  {template.branding.tagline}")
        tag_run.italic = True
        tag_run.font.size = Pt(style.font_size_pt - 1)

    if document.project_name:
        proj_p = doc.add_paragraph()
        proj_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        proj_run = proj_p.add_run(document.project_name)
        proj_run.bold = True

    for content in document.sections:
        _render_section(doc, content, style=size_bundle)

    if template.branding.confidential_footer:
        # The footer renders outside the composer's token pass, so resolve [COMPANY] here.
        footer_text = template.branding.confidential_footer.replace(
            "[COMPANY]", template.branding.company
        )
        footer_p = doc.sections[0].footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_p.add_run(footer_text)
        footer_run.italic = True
        footer_run.font.size = Pt(style.font_size_pt - 2)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
